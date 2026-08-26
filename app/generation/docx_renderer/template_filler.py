"""
DOCX CV renderer.

Loads the supervisor-provided Word template and fills it using the
cv_json produced by the existing CV matching/generation pipeline.

The renderer deliberately works on the existing template instead of
creating a new Word document from scratch.
"""

from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[3]

TEMPLATE_PATH = PROJECT_ROOT / "Templates" / "Template_CV_DVT.docx"


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------

def _text(value) -> str:
    """Safely convert a value to displayable text."""
    if value is None:
        return ""
    return str(value).strip()


def _clear_paragraph(paragraph) -> None:
    """Remove all runs from a paragraph."""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def _set_paragraph_text(paragraph, text: str, bold: bool = False) -> None:
    """Replace paragraph content while keeping the paragraph itself."""
    _clear_paragraph(paragraph)

    run = paragraph.add_run(_text(text))
    run.bold = bold


def _clear_cell(cell) -> None:
    """Clear all paragraphs from a table cell."""
    cell.text = ""


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    """Replace a cell's content."""
    _clear_cell(cell)

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(_text(text))
    run.bold = bold


def _add_bullets(cell, items) -> None:
    """Write a list as bullet paragraphs inside a table cell."""
    _clear_cell(cell)

    if not items:
        return

    first = True

    for item in items:
        if isinstance(item, dict):
            value = (
                item.get("description")
                or item.get("category")
                or item.get("name")
                or ""
            )
        else:
            value = item

        value = _text(value)

        if not value:
            continue

        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False

        paragraph.add_run(f"• {value}")


def _format_experience_bullets(exp: dict) -> list[str]:
    """
    Convert an experience into displayable activity bullets.

    Uses responsibilities first, then deliverables and description.
    """

    result = []

    # Responsibilities
    responsibilities = exp.get("responsibilities") or []

    for responsibility in responsibilities:
        if isinstance(responsibility, dict):
            value = (
                responsibility.get("description")
                or responsibility.get("category")
                or ""
            )
        else:
            value = responsibility

        value = _text(value)

        if value and value not in result:
            result.append(value)

    # Deliverables
    deliverables = exp.get("deliverables") or []

    for deliverable in deliverables:
        value = _text(deliverable)

        if value and value not in result:
            result.append(value)

    # If there are no responsibilities/deliverables,
    # use the experience description.
    if not result:
        description = _text(exp.get("description"))

        if description:
            result.append(description)

    return result
def _experience_company_title(exp: dict) -> str:
    company = _text(exp.get("company"))
    title = _text(exp.get("title"))

    if company and title:
        return f"{company} - {title}"

    return company or title


def _experience_period(exp: dict) -> str:
    """
    Get the experience period from the different structures
    used by the CV pipeline.
    """

    # Yasmine's CV JSON uses "dates".
    dates = _text(exp.get("dates"))

    if dates and dates != "Not specified":
        return dates

    if exp.get("period"):
        return _text(exp["period"])

    start = (
        exp.get("start_date")
        or exp.get("start")
        or exp.get("from")
        or ""
    )

    end = (
        exp.get("end_date")
        or exp.get("end")
        or exp.get("to")
        or ""
    )

    start = _text(start)
    end = _text(end)

    if start and end:
        return f"{start} - {end}"

    if start:
        return f"{start} - Présent"

    return ""
# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _find_table_with_keywords(document, keywords):
    """
    Find the first table containing one of the requested keywords.

    This is intentionally tolerant because Word templates can move slightly
    between versions.
    """

    keywords = [k.lower() for k in keywords]

    for table in document.tables:
        table_text = " ".join(
            cell.text.lower()
            for row in table.rows
            for cell in row.cells
        )

        if any(keyword in table_text for keyword in keywords):
            return table

    return None


def _find_paragraph_with_text(document, text):
    """Find a paragraph containing the requested text."""
    text = text.lower()

    for paragraph in document.paragraphs:
        if text in paragraph.text.lower():
            return paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if text in paragraph.text.lower():
                        return paragraph

    return None
def _replace_table_cell_value(document: Document, label: str, value) -> bool:
    """
    Find a table cell containing `label` and replace the value
    in the cell immediately to its right.
    """
    value = _text(value)

    if not value:
        return False

    label = label.lower().strip()

    for table in document.tables:
        for row in table.rows:
            cells = row.cells

            for i, cell in enumerate(cells):
                cell_text = cell.text.lower().strip()

                if label in cell_text:
                    if i + 1 < len(cells):
                        _set_cell_text(
                            cells[i + 1],
                            value,
                        )
                        return True

    return False

# ---------------------------------------------------------------------------
# Identity / header
# ---------------------------------------------------------------------------

def _replace_table_cell_value(document: Document, label: str, value) -> bool:
    """
    Find a table cell containing `label` and replace the value
    in the cell immediately to its right.
    """
    value = _text(value)

    if not value:
        return False

    label = label.lower().strip()

    for table in document.tables:
        for row in table.rows:
            cells = row.cells

            for i, cell in enumerate(cells):
                cell_text = cell.text.lower().strip()

                if label in cell_text:
                    if i + 1 < len(cells):
                        _set_cell_text(
                            cells[i + 1],
                            value,
                        )
                        return True

    return False


def fill_identity(document: Document, cv_json: dict) -> None:
    """
    Fill the identity/header table of the supervisor template.

    The template stores identity information in the first table:
        row 0 = Title
        row 1 = Name
        row 2 = Date of birth
        row 3 = Nationality / residence
    """

    if not document.tables:
        return

    table = document.tables[0]

    if len(table.rows) < 4:
        return

    # Row 0: Title
    title = _text(cv_json.get("title"))
    if title and len(table.rows[0].cells) >= 2:
        table.rows[0].cells[1].text = title

    # Row 1: Name
    name = _text(cv_json.get("name"))
    if name and len(table.rows[1].cells) >= 2:
        table.rows[1].cells[1].text = name

    # Row 2: Date of birth
    date_of_birth = _text(cv_json.get("date_of_birth"))
    if date_of_birth and len(table.rows[2].cells) >= 2:
        table.rows[2].cells[1].text = date_of_birth
    else:
        # Remove the template candidate's value
        table.rows[2].cells[1].text = ""

    # Row 3: Nationality / residence
    nationality = _text(cv_json.get("nationality"))
    if nationality and len(table.rows[3].cells) >= 2:
        table.rows[3].cells[1].text = nationality
    else:
        # Remove the template candidate's value
        table.rows[3].cells[1].text = ""
# ---------------------------------------------------------------------------
# Education / training
# ---------------------------------------------------------------------------

def fill_education(document: Document, cv_json: dict) -> None:
    education = cv_json.get("education") or []

    paragraph = _find_paragraph_with_text(document, "Etudes")

    if paragraph is None:
        return

    # Find the "Autres Formations / Certifications" paragraph.
    next_section = _find_paragraph_with_text(
        document,
        "Autres Formations / Certifications",
    )

    # Remove the old education paragraphs between:
    # "Etudes :" and "Autres Formations / Certifications :"
    current = paragraph._p.getnext()

    while current is not None and (
        next_section is None or current is not next_section._p
    ):
        next_element = current.getnext()
        current.getparent().remove(current)
        current = next_element

    # Nothing new to insert.
    if not education:
        return

    # Insert the new education paragraphs immediately after "Etudes :".
    previous = paragraph._p

    for item in education:
        if not isinstance(item, dict):
            text = _text(item)
        else:
            parts = [
                item.get("degree"),
                item.get("field_of_study"),
                item.get("institution"),
            ]

            text = " - ".join(
                _text(part)
                for part in parts
                if _text(part)
            )

            years = _text(item.get("years"))

            if years:
                text = f"{text} ({years})" if text else years

        if not text:
            continue

        # Create a copy of the original Etudes paragraph
        # so the formatting remains compatible with the template.
        new_p = deepcopy(paragraph._p)

        # Clear its existing text.
        for child in list(new_p):
            if child.tag.endswith("}r"):
                new_p.remove(child)

        previous.addnext(new_p)

        from docx.text.paragraph import Paragraph

        new_paragraph = Paragraph(new_p, paragraph._parent)
        new_paragraph.text = f"• {text}"

        previous = new_p
def fill_other_training(document: Document, cv_json: dict) -> None:
    certifications = cv_json.get("certifications") or []

    paragraph = _find_paragraph_with_text(
        document,
        "Autres Formations / Certifications",
    )

    if paragraph is None:
        return

    # Find the next major section.
    next_section = _find_paragraph_with_text(
        document,
        "Expérience professionnelle pertinente",
    )

    # Remove the old certifications/training paragraphs.
    current = paragraph._p.getnext()

    while current is not None and (
        next_section is None or current is not next_section._p
    ):
        next_element = current.getnext()
        current.getparent().remove(current)
        current = next_element

    if not certifications:
        return

    # Insert the new certifications immediately after the heading.
    previous = paragraph._p

    from docx.text.paragraph import Paragraph

    for cert in certifications:
        if isinstance(cert, dict):
            name = cert.get("name") or cert.get("title") or ""
            issuer = cert.get("issuer") or ""
            year = cert.get("year") or ""

            parts = [
                _text(name),
                _text(issuer),
                _text(year),
            ]

            value = " - ".join(
                part for part in parts if part
            )
        else:
            value = cert

        value = _text(value)

        if not value:
            continue

        # Copy the heading paragraph to preserve template formatting.
        new_p = deepcopy(paragraph._p)

        # Remove the existing text/runs from the copy.
        for child in list(new_p):
            if child.tag.endswith("}r"):
                new_p.remove(child)

        previous.addnext(new_p)

        new_paragraph = Paragraph(
            new_p,
            paragraph._parent,
        )

        new_paragraph.text = f"• {value}"

        previous = new_p

# ---------------------------------------------------------------------------
# Languages
# ---------------------------------------------------------------------------

def fill_languages(document: Document, cv_json: dict) -> None:
    languages = cv_json.get("languages") or []

    paragraph = _find_paragraph_with_text(
        document,
        "Langues pratiquées",
    )

    if paragraph is None:
        return

    # Find the next major section.
    next_section = _find_paragraph_with_text(
        document,
        "Compétences/qualifications",
    )

    # Remove the old language paragraphs from the template.
    current = paragraph._p.getnext()

    while current is not None and (
        next_section is None or current is not next_section._p
    ):
        next_element = current.getnext()
        current.getparent().remove(current)
        current = next_element

    if not languages:
        return

    from docx.text.paragraph import Paragraph

    previous = paragraph._p

    for language in languages:
        if isinstance(language, dict):
            name = (
                language.get("language")
                or language.get("name")
                or ""
            )

            level = (
                language.get("level")
                or language.get("proficiency")
                or ""
            )

            value = f"{name} : {level}" if level else name
        else:
            value = language

        value = _text(value)

        if not value:
            continue

        # Copy the heading paragraph to preserve template formatting.
        new_p = deepcopy(paragraph._p)

        # Remove the existing text/runs from the copied paragraph.
        for child in list(new_p):
            if child.tag.endswith("}r"):
                new_p.remove(child)

        previous.addnext(new_p)

        new_paragraph = Paragraph(
            new_p,
            paragraph._parent,
        )

        new_paragraph.text = f"• {value}"

        previous = new_p
# ---------------------------------------------------------------------------
# Experience tables
# ---------------------------------------------------------------------------

def fill_experience_table(table, experiences: list[dict]) -> None:
    """
    Fill an existing 4-column experience table.

    Expected columns:

        0 = Period
        1 = Employer / title / references
        2 = Country
        3 = Activities
    """

    if table is None:
        return

    if not experiences:
        return

    # Keep the header row and remove existing data rows.
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for exp in experiences:
        row = table.add_row()

        cells = row.cells

        if len(cells) < 4:
            continue

        _set_cell_text(
            cells[0],
            _experience_period(exp),
        )

        _set_cell_text(
            cells[1],
            _experience_company_title(exp),
        )

        _set_cell_text(
            cells[2],
            _text(
                exp.get("country")
                or exp.get("location")
            ),
        )

        _add_bullets(
            cells[3],
            _format_experience_bullets(exp),
        )


def fill_professional_experience(document: Document, cv_json: dict) -> None:
    experiences = cv_json.get("experience") or []

    if not experiences:
        return

    table = _find_table_with_keywords(
        document,
        [
            "nom de l’employeur",
            "nom de l'employeur",
            "sommaire des activités",
            "activités réalisées",
        ],
    )

    if table is None:
        return

    fill_experience_table(table, experiences)


def fill_mission_experience(document: Document, cv_json: dict) -> None:
    """
    Fill the mission-specific experience section.

    The existing pipeline calls these objects `projects`, so they are mapped
    into the same 4-column Word structure.
    """

    projects = cv_json.get("projects") or []

    if not projects:
        return

    tables = document.tables

    if len(tables) < 2:
        return

    # In the supplied template, the mission-specific table occurs after
    # the general professional experience content.
    candidate_tables = [
        table
        for table in tables
        if "sommaire des activités" in " ".join(
            cell.text.lower()
            for row in table.rows
            for cell in row.cells
        )
    ]

    if len(candidate_tables) < 2:
        return

    mission_table = candidate_tables[-1]

    fill_experience_table(
        mission_table,
        projects,
    )


# ---------------------------------------------------------------------------
# Contact
# ---------------------------------------------------------------------------

def fill_contact(document: Document, cv_json: dict) -> None:
    email = _text(cv_json.get("email"))
    phone = _text(cv_json.get("phone"))

    paragraph = _find_paragraph_with_text(
        document,
        "Renseignements pour contacter",
    )

    if paragraph is None:
        return

    # Remove ALL existing content from the paragraph.
    # This removes the old template email and phone as well.
    for child in list(paragraph._p):
        # Keep paragraph properties (<w:pPr>) if present.
        if child.tag.endswith("}pPr"):
            continue

        paragraph._p.remove(child)

    # Build the new contact information.
    parts = []

    if email:
        parts.append(f"courriel : {email}")

    if phone:
        parts.append(f"téléphone : {phone}")

    text = "Renseignements pour contacter l’expert :"

    if parts:
        text += " " + ", ".join(parts)

    paragraph.add_run(text)
def fill_signature(document: Document, cv_json: dict) -> None:
    name = _text(cv_json.get("name"))

    if not name:
        return

    for paragraph in document.paragraphs:
        if "Signature" in paragraph.text and "Date" in paragraph.text:
            # Preserve the template structure.
            paragraph.text = f"{name}\tSignature \tDate"
            return
# ---------------------------------------------------------------------------
# Main renderer
# ---------------------------------------------------------------------------

def render_cv_docx(
    cv_json: dict,
    output_path: str,
    target_language: str = "French",
) -> str:
    """
    Generate a DOCX CV from the supervisor's Word template.

    Parameters
    ----------
    cv_json:
        Existing CV JSON produced by the project's matching pipeline.

    output_path:
        Destination .docx file.

    target_language:
        Language selected for generation. Currently the renderer preserves
        the language already present in cv_json.

    Returns
    -------
    str
        Generated DOCX path.
    """

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"DOCX template not found: {TEMPLATE_PATH}"
        )

    output = Path(output_path)

    output.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document(str(TEMPLATE_PATH))

    fill_identity(document, cv_json)
    fill_education(document, cv_json)
    fill_other_training(document, cv_json)
    fill_languages(document, cv_json)
    fill_professional_experience(document, cv_json)
    fill_mission_experience(document, cv_json)
    fill_contact(document, cv_json)

    document.save(str(output))

    return str(output)